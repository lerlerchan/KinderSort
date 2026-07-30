"""
export_matched_photos.py — Export ONLY matched photos from Output folder to Word

This script exports photos that were successfully matched to students.
Use this when you have many matched photos and want to organize them.

Usage:
    python export_matched_photos.py
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
import shutil
import os


# ============= CONFIGURATION =============
# Modify these settings before running

OUTPUT_FOLDER = Path('Output')  # KinderSort output folder
OUTPUT_WORD_FILE = Path('KinderSort_Matched_Photos_Report.docx')

# Filtering Options
MIN_PHOTOS_PER_STUDENT = 1  # Minimum photos to include a student (set to 0 for all)
MAX_PHOTOS_PER_STUDENT = 50  # Maximum photos per student
INCLUDE_UNMATCHED = False  # Set to True to include unmatched photos section

# Export specific students only (leave as None for all)
# Example: SPECIFIC_STUDENTS = ['Amelie Mauresmo', 'Clint Eastwood']
SPECIFIC_STUDENTS = None

# Thumbnail Settings
THUMBNAIL_WIDTH = 250  # pixels
THUMBNAIL_HEIGHT = 250  # pixels
COLS_PER_ROW = 3  # Number of photos per row

# =========================================


def setup_temp_folder():
    """Create temporary folder for thumbnails"""
    temp_dir = Path('temp_thumbnails')
    temp_dir.mkdir(exist_ok=True)
    return temp_dir


def cleanup_temp_folder(temp_dir: Path):
    """Remove temporary thumbnail folder"""
    if temp_dir.exists():
        shutil.rmtree(temp_dir)


def resize_image_for_word(img_path: Path, temp_dir: Path, max_width: int = 250, max_height: int = 250) -> Path:
    """Resize image for Word document thumbnail"""
    # Create safe filename
    safe_name = ''.join(c for c in img_path.stem if c.isalnum() or c in '._-')
    safe_name = safe_name[:40]
    temp_path = temp_dir / f"thumb_{safe_name}{img_path.suffix}"
    
    try:
        with Image.open(img_path) as img:
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            img.thumbnail((max_width, max_height), Image.LANCZOS)
            img.save(temp_path, 'JPEG', quality=85)
            return temp_path
    except Exception as e:
        print(f"  ⚠️ Could not resize {img_path.name}: {e}")
        return img_path


def get_image_files(folder: Path) -> list[Path]:
    """Get all image files in a folder"""
    extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.webp', '.gif'}
    return sorted([f for f in folder.iterdir() if f.is_file() and f.suffix.lower() in extensions])


def get_student_data(output_folder: Path, min_photos: int = 1, specific_students: list = None) -> list[tuple]:
    """Get student folders with photo counts"""
    student_data = []
    
    for folder in output_folder.iterdir():
        if not folder.is_dir() or folder.name == '_unmatched':
            continue
        
        if specific_students and folder.name not in specific_students:
            continue
        
        photos = get_image_files(folder)
        count = len(photos)
        
        if count >= min_photos:
            student_data.append((folder, folder.name, count))
    
    student_data.sort(key=lambda x: x[2], reverse=True)
    return student_data


def export_matched_photos_only(
    output_folder: Path,
    output_path: Path,
    min_photos: int = 1,
    max_photos: int = 50,
    include_unmatched: bool = False,
    specific_students: list = None,
    cols: int = 3
) -> None:
    """Export only matched photos to Word document"""
    
    print(f"📂 Scanning output folder: {output_folder}")
    
    # Get student data
    student_data = get_student_data(output_folder, min_photos, specific_students)
    
    if not student_data:
        print(f"❌ No students found with {min_photos}+ photos!")
        print("   Try setting MIN_PHOTOS_PER_STUDENT = 0")
        return
    
    total_photos = sum(count for _, _, count in student_data)
    print(f"   Found {len(student_data)} students with {total_photos} matched photos")
    
    # Get unmatched count
    unmatched_folder = output_folder / '_unmatched'
    unmatched_count = len(get_image_files(unmatched_folder)) if unmatched_folder.exists() else 0
    
    if unmatched_count > 0:
        print(f"   Found {unmatched_count} unmatched photos (not included)")
    
    # Setup temp folder
    temp_dir = setup_temp_folder()
    
    # Create Word document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== Title =====
    title = doc.add_heading('KinderSort — Matched Photos Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== Metadata =====
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Matched Photos: {total_photos}')
    doc.add_paragraph(f'Total Students: {len(student_data)}')
    if specific_students:
        doc.add_paragraph(f'Filtered Students: {", ".join(specific_students)}')
    doc.add_paragraph()
    
    # ===== Summary Table =====
    doc.add_heading('📊 Summary', level=2)
    
    table = doc.add_table(rows=len(student_data) + 2, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    table.rows[0].cells[0].text = 'Student Name'
    table.rows[0].cells[1].text = 'Photo Count'
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student rows
    for i, (_, name, count) in enumerate(student_data, start=1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = str(count)
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Total row
    total_row = len(student_data) + 1
    table.rows[total_row].cells[0].text = 'TOTAL'
    table.rows[total_row].cells[1].text = str(total_photos)
    table.rows[total_row].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[total_row].cells[1].paragraphs[0].runs[0].bold = True
    table.rows[total_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ===== Photos for each student =====
    for folder, student_name, photo_count in student_data:
        images = get_image_files(folder)
        
        if not images:
            continue
        
        doc.add_page_break()
        doc.add_heading(f'👨‍🎓 {student_name} ({len(images)} photos)', level=2)
        doc.add_paragraph()
        
        display_images = images[:max_photos]
        if len(images) > max_photos:
            doc.add_paragraph(f'Showing first {max_photos} of {len(images)} photos')
            doc.add_paragraph()
        
        rows = (len(display_images) + cols - 1) // cols
        
        if rows > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.autofit = False
            
            for col in range(cols):
                table.columns[col].width = Inches(2.0)
            
            row_idx = 0
            col_idx = 0
            
            for img_path in display_images:
                cell = table.cell(row_idx, col_idx)
                
                try:
                    thumb_path = resize_image_for_word(img_path, temp_dir, THUMBNAIL_WIDTH, THUMBNAIL_HEIGHT)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(str(thumb_path), width=Inches(1.8))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    filename_para = cell.add_paragraph()
                    filename_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    filename_run = filename_para.add_run(img_path.name[:25] + ('...' if len(img_path.name) > 25 else ''))
                    filename_run.font.size = Pt(7)
                    filename_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    
                except Exception as e:
                    cell.text = f'⚠️ {img_path.name[:15]}'
                
                col_idx += 1
                if col_idx >= cols:
                    col_idx = 0
                    row_idx += 1
        
        doc.add_paragraph()
    
    # ===== Unmatched Photos (optional) =====
    if include_unmatched:
        unmatched_folder = output_folder / '_unmatched'
        if unmatched_folder.exists():
            unmatched_images = get_image_files(unmatched_folder)
            if unmatched_images:
                doc.add_page_break()
                doc.add_heading('❓ Unmatched Photos', level=2)
                doc.add_paragraph(f'{len(unmatched_images)} photos were not matched:')
                doc.add_paragraph()
                
                display_unmatched = unmatched_images[:30]
                rows = (len(display_unmatched) + cols - 1) // cols
                
                if rows > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    table.autofit = False
                    
                    for col in range(cols):
                        table.columns[col].width = Inches(2.0)
                    
                    row_idx = 0
                    col_idx = 0
                    
                    for img_path in display_unmatched:
                        cell = table.cell(row_idx, col_idx)
                        
                        try:
                            thumb_path = resize_image_for_word(img_path, temp_dir, THUMBNAIL_WIDTH, THUMBNAIL_HEIGHT)
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(str(thumb_path), width=Inches(1.8))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            filename_para = cell.add_paragraph()
                            filename_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            filename_run = filename_para.add_run(img_path.name[:25])
                            filename_run.font.size = Pt(7)
                            filename_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        except Exception:
                            cell.text = '⚠️ Error'
                        
                        col_idx += 1
                        if col_idx >= cols:
                            col_idx = 0
                            row_idx += 1
                    
                    if len(unmatched_images) > 30:
                        doc.add_paragraph(f'... and {len(unmatched_images) - 30} more')
    
    # ===== Footer =====
    doc.add_paragraph()
    doc.add_paragraph('---')
    footer = doc.add_paragraph('Generated by KinderSort Photo Exporter')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # ===== Save =====
    doc.save(str(output_path))
    cleanup_temp_folder(temp_dir)
    
    print(f'\n✅ Report saved: {output_path.resolve()}')
    print(f'   File size: {output_path.stat().st_size:,} bytes')
    print(f'   Students: {len(student_data)}')
    print(f'   Total Matched Photos: {total_photos}')
    if include_unmatched:
        print(f'   Unmatched Photos: {unmatched_count}')


def main():
    """Main entry point"""
    
    # Check if output folder exists
    if not OUTPUT_FOLDER.exists():
        print(f'❌ Error: {OUTPUT_FOLDER} does not exist')
        print('   Please run KinderSort first to process photos')
        return
    
    # Check for student folders
    student_folders = [f for f in OUTPUT_FOLDER.iterdir() if f.is_dir() and f.name != '_unmatched']
    if not student_folders:
        print(f'⚠️ No student folders found in {OUTPUT_FOLDER}')
        print('   Please run KinderSort first to process photos')
        return
    
    # Display configuration
    print('\n' + '='*60)
    print('📸 KinderSort - Export Matched Photos Only')
    print('='*60)
    print(f'📂 Output Folder: {OUTPUT_FOLDER}')
    print(f'📄 Output File: {OUTPUT_WORD_FILE}')
    print(f'👨‍🎓 Min Photos per Student: {MIN_PHOTOS_PER_STUDENT}')
    print(f'📸 Max Photos per Student: {MAX_PHOTOS_PER_STUDENT}')
    print(f'❓ Include Unmatched: {INCLUDE_UNMATCHED}')
    if SPECIFIC_STUDENTS:
        print(f'🎯 Specific Students: {", ".join(SPECIFIC_STUDENTS)}')
    print('='*60)
    print()
    
    try:
        export_matched_photos_only(
            output_folder=OUTPUT_FOLDER,
            output_path=OUTPUT_WORD_FILE,
            min_photos=MIN_PHOTOS_PER_STUDENT,
            max_photos=MAX_PHOTOS_PER_STUDENT,
            include_unmatched=INCLUDE_UNMATCHED,
            specific_students=SPECIFIC_STUDENTS,
            cols=COLS_PER_ROW
        )
        print('\n🎉 Done! Open the Word document to view all matched photos.')
    except Exception as e:
        print(f'\n❌ Error: {e}')
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()