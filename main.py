"""
main.py — KinderSort GUI entry point.

Single-window tkinter application that drives the PhotoSorter pipeline with a
background thread so the UI remains responsive during processing.
"""

import threading
import time
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

from sorter import PhotoSorter
from utils import setup_logger

# Check if python-docx is installed (for generating student Word reports)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ---------------------------------------------------------------------------
# NVIDIA API Integration and Initialization Configuration
# ---------------------------------------------------------------------------
BUILTIN_NVIDIA_API_KEY = os.environ.get(
    "NVIDIA_API_KEY", 
    "YOUR_API_KEY"  
)



class KinderSortApp(tk.Tk):
    """Main application window for KinderSort — Student Photo Organiser."""

    MIN_WIDTH = 500
    MIN_HEIGHT = 400

    def __init__(self) -> None:
        """Initialise the window, build all widgets, and configure layout."""
        super().__init__()
        self.title("KinderSort v1.1 — Student Photo Organiser")
        self.minsize(self.MIN_WIDTH, self.MIN_HEIGHT)
        self.resizable(True, True)

        # StringVars for the three folder paths
        self._reference_var = tk.StringVar()
        self._events_var = tk.StringVar()
        self._output_var = tk.StringVar()

        # Cancellation flag shared between GUI and worker thread
        self._cancel_flag = threading.Event()

        # Spinner / elapsed timer state
        self._spinner_frames = ["🕐", "🕑", "🕒", "🕓", "🕔", "🕕", "🕖", "🕗", "🕘", "🕙", "🕚", "🕛"]
        self._spinner_idx = 0
        self._sort_start_time: float | None = None
        self._ticker_id: str | None = None

        self._build_ui()

    # ------------------------------------------------------------------
    # Widget construction
    # ------------------------------------------------------------------

    def _build_ui(self) -> None:
        """Build and pack all widgets into the main window."""
        root_frame = tk.Frame(self, padx=16, pady=16)
        root_frame.pack(fill=tk.BOTH, expand=True)

        # Title label
        tk.Label(
            root_frame,
            text="KinderSort — Student Photo Organiser",
            font=("Helvetica", 14, "bold"),
        ).pack(anchor="w", pady=(0, 12))

        # Folder selector rows
        folders_frame = tk.LabelFrame(root_frame, text="Folders", padx=8, pady=8)
        folders_frame.pack(fill=tk.X, pady=(0, 12))

        self._build_folder_row(folders_frame, "Reference Photos:", self._reference_var, 0)
        self._build_folder_row(folders_frame, "Events Folder:", self._events_var, 1)
        self._build_folder_row(folders_frame, "Output Folder:", self._output_var, 2)

        folders_frame.columnconfigure(1, weight=1)

        # Start / Cancel buttons
        btn_frame = tk.Frame(root_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 12))

        self._start_btn = tk.Button(
            btn_frame,
            text="Start Sorting",
            font=("Helvetica", 11, "bold"),
            bg="#4CAF50",
            fg="white",
            activebackground="#388E3C",
            activeforeground="white",
            padx=16,
            pady=8,
            command=self._on_start,
        )
        self._start_btn.pack(side=tk.LEFT, padx=(0, 8))

        self._cancel_btn = tk.Button(
            btn_frame,
            text="Cancel",
            font=("Helvetica", 11),
            padx=16,
            pady=8,
            state=tk.DISABLED,
            command=self._on_cancel,
        )
        self._cancel_btn.pack(side=tk.LEFT, padx=(0, 8))

        # Optional Feature Button: Export Word Report
        self._export_btn = tk.Button(
            btn_frame,
            text="Generate Roster Docx",
            font=("Helvetica", 10),
            padx=10,
            pady=8,
            command=self._on_export_docx,
        )
        self._export_btn.pack(side=tk.RIGHT)

        # Progress section
        self._build_progress_section(root_frame)

        # Summary box
        self._build_summary_box(root_frame)

    def _build_folder_row(
        self,
        parent: tk.Widget,
        label_text: str,
        string_var: tk.StringVar,
        row: int,
    ) -> None:
        """Create a label + read-only entry + browse button row inside parent.

        Args:
            parent: Container widget (expects grid layout).
            label_text: Text displayed on the left label.
            string_var: StringVar bound to the entry widget.
            row: Grid row index.
        """
        tk.Label(parent, text=label_text, anchor="w").grid(
            row=row, column=0, sticky="w", padx=(0, 8), pady=4
        )

        entry = tk.Entry(parent, textvariable=string_var, state="readonly", width=40)
        entry.grid(row=row, column=1, sticky="ew", pady=4)

        btn = tk.Button(
            parent,
            text="Browse…",
            command=lambda v=string_var: self._browse_folder(v),
        )
        btn.grid(row=row, column=2, padx=(8, 0), pady=4)

    def _build_progress_section(self, parent: tk.Widget) -> None:
        """Build the progress bar and status label."""
        progress_frame = tk.LabelFrame(parent, text="Progress", padx=8, pady=8)
        progress_frame.pack(fill=tk.X, pady=(0, 12))

        self._progress_var = tk.DoubleVar(value=0.0)
        self._progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self._progress_var,
            maximum=100,
            mode="determinate",
        )
        self._progress_bar.pack(fill=tk.X, pady=(0, 4))

        self._status_label = tk.Label(
            progress_frame, text="Ready.", anchor="w", wraplength=460
        )
        self._status_label.pack(fill=tk.X)

        self._timer_label = tk.Label(
            progress_frame, text="", anchor="w", fg="#555555"
        )
        self._timer_label.pack(fill=tk.X)

    def _build_summary_box(self, parent: tk.Widget) -> None:
        """Build the read-only summary text box shown after completion."""
        summary_frame = tk.LabelFrame(parent, text="Summary", padx=8, pady=8)
        summary_frame.pack(fill=tk.BOTH, expand=True)

        self._summary_text = tk.Text(
            summary_frame, height=5, state=tk.DISABLED, wrap=tk.WORD
        )
        self._summary_text.pack(fill=tk.BOTH, expand=True)

    # ------------------------------------------------------------------
    # Event handlers
    # ------------------------------------------------------------------

    def _browse_folder(self, string_var: tk.StringVar) -> None:
        """Open a directory chooser and update string_var with the selection."""
        folder = filedialog.askdirectory(title="Select folder")
        if folder:
            string_var.set(folder)

    def _on_start(self) -> None:
        """Validate inputs then launch the worker thread for all heavy work."""
        ref = self._reference_var.get().strip().strip('"')
        events = self._events_var.get().strip().strip('"')
        output = self._output_var.get().strip().strip('"')

        if not ref or not events or not output:
            messagebox.showerror(
                "Missing folders",
                "Please select all three folders before starting.",
            )
            return

        ref_path = Path(ref)
        events_path = Path(events)
        output_path = Path(output)

        for path, name in [(ref_path, "Reference"), (events_path, "Events")]:
            if not path.is_dir():
                messagebox.showerror("Invalid folder", f"{name} folder does not exist:\n{path}")
                return

        # Ensure output folder is creatable / writable
        try:
            output_path.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            messagebox.showerror("Output folder error", f"Cannot create output folder:\n{exc}")
            return

        # Disable start, enable cancel before launching thread
        self._start_btn.config(state=tk.DISABLED)
        self._cancel_btn.config(state=tk.NORMAL)
        self._cancel_flag.clear()
        self._clear_summary()
        self._progress_var.set(0)
        self._set_status("Loading reference photos…")
        self._start_ticker()

        logger = setup_logger(output_path)

        sorter = PhotoSorter(
            ref_path, 
            events_path, 
            output_path, 
            logger, 
            api_key=BUILTIN_NVIDIA_API_KEY
        )

        thread = threading.Thread(
            target=self._run_sorting, args=(sorter,), daemon=True
        )
        thread.start()

    def _run_sorting(self, sorter: PhotoSorter) -> None:
        """Worker thread: load references, then sort all photos."""
        try:
            skipped_names = sorter.load_references(
                progress_callback=self._on_ref_progress
            )
        except FileNotFoundError as exc:
            self.after(0, self._on_error, f"File path error: {exc}")
            return
        except PermissionError as exc:
            self.after(0, self._on_error, f"Permission denied accessing folder: {exc}")
            return
        except Exception as exc:  # noqa: BLE001
            self.after(0, self._on_error, str(exc))
            return

        if skipped_names:
            self.after(0, self._show_ref_warning, skipped_names)

        if not sorter._student_encodings:
            self.after(0, self._on_error, "No student faces could be loaded. Please check your Reference folder.")
            return

        try:
            summary = sorter.sort_all(
                progress_callback=self._on_progress,
                cancelled=self._cancel_flag.is_set,
            )
            self.after(0, self._on_done, summary)
        except FileNotFoundError as exc:
            self.after(0, self._on_error, f"File path error: {exc}")
            return
        except PermissionError as exc:
            self.after(0, self._on_error, f"Permission denied accessing folder: {exc}")
            return
        except Exception as exc:  # noqa: BLE001
            self.after(0, self._on_error, str(exc))
            return

    def _start_ticker(self) -> None:
        """Start the spinning clock emoji and elapsed timer."""
        self._sort_start_time = time.monotonic()
        self._spinner_idx = 0
        self._tick()

    def _tick(self) -> None:
        """Update spinner and elapsed time every 250 ms."""
        if self._sort_start_time is None:
            return
        elapsed = int(time.monotonic() - self._sort_start_time)
        minutes, seconds = divmod(elapsed, 60)
        spinner = self._spinner_frames[self._spinner_idx % len(self._spinner_frames)]
        self._spinner_idx += 1
        self._timer_label.config(text=f"{spinner}  {minutes:02d}:{seconds:02d} elapsed")
        self._ticker_id = self.after(250, self._tick)

    def _stop_ticker(self, final_elapsed: int | None = None) -> None:
        """Stop the spinner and show final elapsed time."""
        if self._ticker_id:
            self.after_cancel(self._ticker_id)
            self._ticker_id = None
        if final_elapsed is not None:
            minutes, seconds = divmod(final_elapsed, 60)
            self._timer_label.config(text=f"✅  Done in {minutes:02d}:{seconds:02d}")
        else:
            self._timer_label.config(text="")
        self._sort_start_time = None

    def _on_ref_progress(self, current: int, total: int, name: str) -> None:
        """Called from worker thread after each reference photo is encoded."""
        self.after(0, self._set_status, f"Loading references [{current}/{total}]: {name}…")

    def _show_ref_warning(self, skipped_names: list[str]) -> None:
        """Show warning dialog for reference photos with no detectable face."""
        names_str = "\n".join(f"  • {n}" for n in skipped_names)
        messagebox.showwarning(
            "Reference photos without faces",
            f"No face was detected in the reference photos for:\n\n{names_str}\n\n"
            "These students will be skipped during sorting.",
        )

    def _on_cancel(self) -> None:
        """Signal the worker thread to stop after the current image."""
        self._cancel_flag.set()
        self._cancel_btn.config(state=tk.DISABLED)
        self._set_status("Cancelling… (finishing current image)")

     # ------------------------------------------------------------------
    # Optional Feature: Word Grid Report Generator
    # ------------------------------------------------------------------
    def _on_export_docx(self) -> None:
        """Generate a 1-page Word document grid for reference photos."""
        if not HAS_DOCX:
            messagebox.showerror(
                "Missing Library",
                "python-docx is required for this feature.\nPlease run: pip install python-docx",
            )
            return

        ref_path = Path(ref)
        out_path = Path(out_dir) if out_dir else ref_path
        out_path.mkdir(parents=True, exist_ok=True)

        if not ref_dir:
            messagebox.showerror("Error", "Please select a Reference Photos folder first.")
            return

        ref_path = Path(ref_dir)
        out_path = Path(out_dir) if out_dir else ref_path

        valid_exts = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
        images = [p for p in ref_path.iterdir() if p.suffix.lower() in valid_exts]

        if not images:
            messagebox.showwarning("No Images", "No reference images found in the selected folder.")
            return

        try:
            doc = Document()
            # Set margins to 0.5 inch for 1-page fit
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            title = doc.add_heading("Kindergarten Student Roster Grid", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Create a 3-column table
            cols = 3
            rows = (len(images) + cols - 1) // cols
            table = doc.add_table(rows=rows, cols=cols)
            table.autofit = False

            for idx, img_p in enumerate(sorted(images)):
                r = idx // cols
                c = idx % cols
                cell = table.cell(r, c)
                
                # Add image & student name
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_p), width=Inches(1.8))
                
                name_p = cell.add_paragraph(img_p.stem)
                name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            save_file = out_path / "Student_Roster_Grid.docx"
            doc.save(str(save_file))
            messagebox.showinfo("Success", f"Word Report exported to:\n{save_file}")

        except Exception as exc:
            messagebox.showerror("Report Error", f"Failed to generate Word report:\n{exc}")

    # ------------------------------------------------------------------
    # Cross-thread callbacks (all scheduled via after() from worker)
    # ------------------------------------------------------------------

    def _on_progress(self, current: int, total: int, filename: str) -> None:
        """Update progress bar and status label — called from worker thread via after()."""
        self.after(0, self._apply_progress, current, total, filename)

    def _apply_progress(self, current: int, total: int, filename: str) -> None:
        """Apply progress update on main thread."""
        pct = (current / total * 100) if total else 0
        self._progress_var.set(pct)
        self._set_status(f"[{current}/{total}] {filename}")

    def _on_done(self, summary: dict[str, int]) -> None:
        """Show summary and re-enable controls after successful completion."""
        elapsed = int(time.monotonic() - self._sort_start_time) if self._sort_start_time else None
        self._stop_ticker(final_elapsed=elapsed)
        self._start_btn.config(state=tk.NORMAL)
        self._cancel_btn.config(state=tk.DISABLED)
        self._progress_var.set(100)

        cancelled = self._cancel_flag.is_set()
        status = "Sorting cancelled." if cancelled else "Sorting complete."
        self._set_status(status)

        lines = [
            status,
            "",
            f"Total images found : {summary['total']}",
            f"Matched (sorted)   : {summary['matched']}",
            f"Unmatched          : {summary['unmatched']}",
            f"Skipped (errors)   : {summary['skipped']}",
        ]
        self._write_summary("\n".join(lines))

        #Write runtime summary to the log
        logger = setup_logger(Path(self._output_var.get()))
        logger.info(f"Sorting session finished. Total: {summary['total']}, Matched: {summary['matched']}, Unmatched: {summary['unmatched']}")

        if summary["total"] == 0:
            messagebox.showwarning(
                "No images found",
                "No photos were found in the Events folder.\n\n"
                "Make sure your Events folder contains photos (or sub-folders with photos).\n"
                "Supported formats: .jpg  .jpeg  .png  .bmp  .webp",
            )

    def _on_error(self, message: str) -> None:
        """Show an error dialog and re-enable controls."""
        self._stop_ticker()
        self._start_btn.config(state=tk.NORMAL)
        self._cancel_btn.config(state=tk.DISABLED)
        self._set_status("An error occurred.")
        messagebox.showerror("Unexpected error", message)

    # ------------------------------------------------------------------
    # Widget helpers
    # ------------------------------------------------------------------

    def _set_status(self, text: str) -> None:
        """Update the status label text."""
        self._status_label.config(text=text)

    def _write_summary(self, text: str) -> None:
        """Write text into the read-only summary box."""
        self._summary_text.config(state=tk.NORMAL)
        self._summary_text.delete("1.0", tk.END)
        self._summary_text.insert(tk.END, text)
        self._summary_text.config(state=tk.DISABLED)

    def _clear_summary(self) -> None:
        """Clear the summary text box."""
        self._write_summary("")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    """Launch the KinderSort GUI application."""
    app = KinderSortApp()
    app.mainloop()


if __name__ == "__main__":
    main()
