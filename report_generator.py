import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class StudentReportGenerator:
    """
    KinderSort Lite - Automatic Grid Report Generator (Optional Bonus Task)
    Author: Member 3
    Description: Scans individual student images and auto-generates a clean 
                 1-page Word (.docx) report in grid format.
    """
    
    def __init__(self, output_doc_path="Student_Photos_Summary.docx"):
        self.output_doc_path = output_doc_path

    def generate_grid_report(self, student_name, image_paths):
        """Generates a 1-page structured Word report with images in a grid."""
        doc = Document()
        
        # 页面标题
        heading = doc.add_heading(f"KinderSort Lite - Student Profile: {student_name}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph(f"Total Detected Photos: {len(image_paths)}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if not image_paths:
            doc.save(self.output_doc_path)
            return self.output_doc_path
            
        # 创建网格表格 (例如 2 列网格)
        cols = 2
        rows = (len(image_paths) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for idx, img_path in enumerate(image_paths):
            r = idx // cols
            c = idx % cols
            cell = table.cell(r, c)
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if os.path.exists(img_path):
                # 插入图片并控制尺寸，确保维持单页布局
                cell_paragraph.add_run().add_picture(img_path, width=Inches(2.5))
                
        doc.save(self.output_doc_path)
        print(f"[SUCCESS] Word report generated: {self.output_doc_path}")
        return self.output_doc_path

if __name__ == "__main__":
    generator = StudentReportGenerator()
    print("[INFO] Member 3 Report Generator Engine Ready.")