"""
generate_report.py — Generates a Word (.docx) report for KinderSort sorted results
with a custom Tkinter dialog featuring an "Open Folder" button.
"""

import os
import platform
import subprocess
import tkinter as tk
from tkinter import ttk
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from utils import SUPPORTED_EXTENSIONS


def open_folder_in_explorer(folder_path: Path) -> None:
    """Opens the system file explorer and navigates to the specified directory."""
    folder_str = str(folder_path.resolve())
    try:
        if platform.system() == "Windows":
            os.startfile(folder_str)
        elif platform.system() == "Darwin":  # macOS
            subprocess.run(["open", folder_str], check=False)
        else:  # Linux
            subprocess.run(["xdg-open", folder_str], check=False)
    except Exception as e:
        print(f"Failed to auto-open folder: {e}")


def show_export_success_dialog(parent_window, file_name: str, folder_path: Path) -> None:
    """Displays a custom modal Tkinter dialog showing export completion and offering
    a direct button to open the target folder.
    """
    destroy_root = False
    if parent_window is None:
        parent_window = tk.Tk()
        parent_window.withdraw()
        destroy_root = True

    dialog = tk.Toplevel(parent_window)
    dialog.title("Export Successful")
    dialog.geometry("420x200")
    dialog.resizable(False, False)
    dialog.grab_set()  # Modal dialog configuration

    # Center dialog on screen
    dialog.update_idletasks()
    x = (dialog.winfo_screenwidth() // 2) - (420 // 2)
    y = (dialog.winfo_screenheight() // 2) - (200 // 2)
    dialog.geometry(f"420x200+{x}+{y}")

    main_frame = ttk.Frame(dialog, padding="15")
    main_frame.pack(fill=tk.BOTH, expand=True)

    # 1. Success Header Label
    title_label = ttk.Label(
        main_frame,
        text="✅ Report Generated Successfully!",
        font=("Segoe UI", 11, "bold"),
        foreground="#2e7d32"
    )
    title_label.pack(pady=(0, 10))

    # 2. File and Path Details Frame
    info_frame = ttk.Frame(main_frame)
    info_frame.pack(fill=tk.X, pady=5)

    ttk.Label(info_frame, text=f"File:  {file_name}", font=("Segoe UI", 9, "bold")).pack(anchor=tk.W, pady=2)
    
    path_str = str(folder_path.resolve())
    ttk.Label(info_frame, text=f"Location:  {path_str}", font=("Segoe UI", 9), foreground="#444444", wraplength=380).pack(anchor=tk.W, pady=2)

    # 3. Action Buttons Frame
    btn_frame = ttk.Frame(main_frame)
    btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(15, 0))

    def on_open_folder():
        open_folder_in_explorer(folder_path)
        dialog.destroy()
        if destroy_root:
            parent_window.destroy()

    def on_ok():
        dialog.destroy()
        if destroy_root:
            parent_window.destroy()

    # Open Folder Button
    open_btn = tk.Button(
        btn_frame, 
        text="📁 Open Folder", 
        font=("Segoe UI", 9, "bold"),
        bg="#e0e0e0",
        command=on_open_folder,
        padx=10,
        pady=3
    )
    open_btn.pack(side=tk.LEFT, padx=(20, 10))

    # OK Button
    ok_btn = tk.Button(
        btn_frame, 
        text="OK", 
        font=("Segoe UI", 9, "bold"),
        width=10,
        bg="#e0e0e0",
        command=on_ok,
        pady=3
    )
    ok_btn.pack(side=tk.RIGHT, padx=(10, 20))

    dialog.protocol("WM_DELETE_WINDOW", on_ok)
    dialog.wait_window()


def generate_student_report(
    output_folder: str, 
    report_path: str = "Student_Report.docx",
    parent_window=None
) -> None:
    """Generates a structured grid-format Word report containing sorted student photos.

    Args:
        output_folder (str): Directory containing sorted student subfolders.
        report_path (str): File path for saving the generated .docx report.
        parent_window (optional): Tkinter window reference for anchoring the dialog.
    """
    doc = Document()

    title = doc.add_heading("Student Photo Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = Path(output_folder)
    student_folders = [
        f for f in output_path.iterdir() if f.is_dir() and f.name != "_unmatched"
    ]

    abs_report_path = Path(report_path).resolve()

    if not student_folders:
        doc.add_paragraph("No student folders found.")
        doc.save(str(abs_report_path))
        show_export_success_dialog(parent_window, abs_report_path.name, abs_report_path.parent)
        return

    # Populate table grids for each student
    for student_folder in student_folders:
        student_name = student_folder.name
        doc.add_heading(f"Student: {student_name}", level=2)

        image_files = [
            f for f in student_folder.iterdir() if f.suffix.lower() in SUPPORTED_EXTENSIONS
        ]

        if not image_files:
            doc.add_paragraph("No images found for this student.")
            continue

        cols = 3
        rows = (len(image_files) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for i, img_path in enumerate(image_files):
            row_idx = i // cols
            col_idx = i % cols
            cell = table.cell(row_idx, col_idx)

            try:
                run = cell.paragraphs[0].add_run()
                run.add_picture(str(img_path), width=Inches(1.5))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                cell.text = f"[Error: {os.path.basename(img_path)}]"

        doc.add_paragraph()  # Separator line

    doc.save(str(abs_report_path))
    
    # Trigger success dialog
    show_export_success_dialog(parent_window, abs_report_path.name, abs_report_path.parent)


if __name__ == "__main__":
    generate_student_report("./output", "Student_Report.docx")