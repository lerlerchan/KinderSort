"""export_matched_photos.py — Export matched photos to Word document"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import shutil

OUTPUT_FOLDER = Path('Output')
OUTPUT_WORD_FILE = Path('KinderSort_Matched_Photos_Report.docx')
THUMBNAIL_WIDTH = 200
THUMBNAIL_HEIGHT = 200
COLS_PER_ROW = 3

def get_image_files(folder: Path) -> list[Path]:
    extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.webp'}
    return sorted([f for f in folder.iterdir() if f.is_file() and f.suffix.lower() in extensions])

def resize_image_for_word(img_path: Path, temp_dir: Path, max_width: int = 200, max_height: int = 200) -> Path | None:
    safe_name = ''.join(c for c in img_path.stem if c.isalnum() or c in '._-')
    safe_name = safe_name[:40]
    temp_path = temp_dir / f"thumb_{safe_name}.jpg"
    try:
        with Image.open(img_path) as img:
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            img.thumbnail((max_width, max_height), Image.LANCZOS)
            img.save(str(temp_path), 'JPEG', quality=85)
            return temp_path
    except Exception:
        return None

def main():
    if not OUTPUT_FOLDER.exists():
        print(f'❌ Error: {OUTPUT_FOLDER} does not exist')
        return

    student_folders = [f for f in OUTPUT_FOLDER.iterdir() if f.is_dir() and f.name != '_unmatched']
    if not student_folders:
        print('⚠️ No student folders found')
        return

    print(f'📸 Exporting matched photos to Word...')

    doc = Document()
    title = doc.add_heading('KinderSort — Matched Photos Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    temp_dir = Path('temp_thumbnails')
    temp_dir.mkdir(exist_ok=True)

    try:
        for folder in student_folders:
            images = get_image_files(folder)
            if not images:
                continue

            doc.add_heading(f'👨‍🎓 {folder.name} ({len(images)} photos)', level=2)
            doc.add_paragraph()

            display_images = images[:50]
            rows = (len(display_images) + COLS_PER_ROW - 1) // COLS_PER_ROW

            if rows > 0:
                table = doc.add_table(rows=rows, cols=COLS_PER_ROW)
                table.style = 'Table Grid'

                row_idx = 0
                col_idx = 0
                for img_path in display_images:
                    cell = table.cell(row_idx, col_idx)
                    try:
                        thumb_path = resize_image_for_word(img_path, temp_dir)
                        if thumb_path:
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(str(thumb_path), width=Inches(1.6))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        cell.text = '⚠️ Error'

                    col_idx += 1
                    if col_idx >= COLS_PER_ROW:
                        col_idx = 0
                        row_idx += 1

    finally:
        if temp_dir.exists():
            shutil.rmtree(temp_dir)

    doc.save(str(OUTPUT_WORD_FILE))
    print(f'✅ Report saved: {OUTPUT_WORD_FILE}')

if __name__ == '__main__':
    main()